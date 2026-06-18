import csv
import os
from pathlib import Path

import requests
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from openpyxl import load_workbook

from src import config


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".csv", ".xlsx", ".url"}

# Tags HTML que representam conteúdo principal — ignorar o resto reduz ruído
_TAGS_CONTEUDO = [
    "article", "main", "section", "div", "p",
    "h1", "h2", "h3", "h4", "h5", "h6",
    "li", "td", "th", "blockquote",
]

# Tags que nunca contêm conteúdo útil
_TAGS_RUIDO = [
    "script", "style", "noscript", "header", "footer",
    "nav", "aside", "form", "button", "svg", "img",
    "figure", "figcaption", "iframe", "input", "select",
]


def _extrair_texto_txt(caminho_arquivo: str) -> Document:
    with open(caminho_arquivo, "r", encoding="utf-8") as arquivo:
        conteudo = arquivo.read().strip()

    return Document(page_content=conteudo, metadata={"source": caminho_arquivo, "type": "txt"})


def _extrair_texto_docx(caminho_arquivo: str) -> Document:
    documento = DocxDocument(caminho_arquivo)
    paragrafos = [paragrafo.text.strip() for paragrafo in documento.paragraphs if paragrafo.text.strip()]
    conteudo = "\n".join(paragrafos)

    return Document(page_content=conteudo, metadata={"source": caminho_arquivo, "type": "docx"})


def _extrair_texto_csv(caminho_arquivo: str) -> Document:
    linhas_formatadas = []
    with open(caminho_arquivo, "r", encoding="utf-8-sig", newline="") as arquivo:
        leitor = csv.reader(arquivo)
        for linha in leitor:
            if any(celula.strip() for celula in linha):
                linhas_formatadas.append(" | ".join(celula.strip() for celula in linha))

    conteudo = "\n".join(linhas_formatadas)
    return Document(page_content=conteudo, metadata={"source": caminho_arquivo, "type": "csv"})


def _extrair_texto_xlsx(caminho_arquivo: str) -> list[Document]:
    """
    Extrai texto de um arquivo XLSX gerando Documents auto-contidos.

    Quando a planilha usa agrupamento na primeira coluna (ex.: dia da semana
    aparece só na primeira linha do grupo e as linhas seguintes têm a célula
    vazia), cada grupo vira um Document separado com título, cabeçalho e nome
    do grupo repetidos.  Isso garante que o text splitter nunca gere um chunk
    sem contexto.
    """
    workbook = load_workbook(caminho_arquivo, data_only=True)
    documentos = []

    for aba in workbook.worksheets:
        linhas_brutas: list[list[str]] = []
        for linha in aba.iter_rows(values_only=True):
            valores = ["" if valor is None else str(valor).strip() for valor in linha]
            if any(valores):
                linhas_brutas.append(valores)

        if not linhas_brutas:
            continue

        # Separar cabeçalho (primeira linha com dados) do corpo
        cabecalho = linhas_brutas[0]
        corpo = linhas_brutas[1:]

        if not corpo:
            # Planilha só com cabeçalho — gerar Document simples
            conteudo = f"Planilha: {aba.title}\n" + " | ".join(cabecalho)
            documentos.append(
                Document(
                    page_content=conteudo,
                    metadata={"source": caminho_arquivo, "sheet": aba.title, "type": "xlsx"},
                )
            )
            continue

        # Detectar se a primeira coluna usa agrupamento (muitas células vazias)
        primeira_col_vazia = sum(1 for linha in corpo if not linha[0])
        usa_agrupamento = len(corpo) > 2 and primeira_col_vazia > len(corpo) * 0.3

        if usa_agrupamento:
            documentos.extend(
                _gerar_docs_agrupados(caminho_arquivo, aba.title, cabecalho, corpo)
            )
        else:
            # Sem agrupamento — comportamento original (um Document por aba)
            linhas_formatadas = [" | ".join(cabecalho)]
            linhas_formatadas.extend(" | ".join(vals) for vals in corpo)
            conteudo = f"Planilha: {aba.title}\n" + "\n".join(linhas_formatadas)
            documentos.append(
                Document(
                    page_content=conteudo,
                    metadata={"source": caminho_arquivo, "sheet": aba.title, "type": "xlsx"},
                )
            )

    return documentos


def _gerar_docs_agrupados(
    caminho_arquivo: str,
    nome_aba: str,
    cabecalho: list[str],
    corpo: list[list[str]],
) -> list[Document]:
    """
    Agrupa as linhas pela primeira coluna e gera um Document por grupo.
    Em cada Document o nome do grupo é preenchido em todas as linhas e o
    cabeçalho é repetido, garantindo contexto completo para o retriever.
    """
    from collections import OrderedDict

    grupos: OrderedDict[str, list[list[str]]] = OrderedDict()
    grupo_atual: str | None = None

    for valores in corpo:
        if valores[0]:
            grupo_atual = valores[0]
        if grupo_atual is None:
            continue
        if grupo_atual not in grupos:
            grupos[grupo_atual] = []
        # Preencher a célula vazia com o nome do grupo para contexto
        valores_completos = valores.copy()
        if not valores_completos[0]:
            valores_completos[0] = grupo_atual
        grupos[grupo_atual].append(valores_completos)

    cabecalho_str = " | ".join(cabecalho)
    documentos = []

    for nome_grupo, linhas_grupo in grupos.items():
        linhas_formatadas = [f"  - {' | '.join(vals)}" for vals in linhas_grupo]
        conteudo = (
            f"Planilha: {nome_aba}\n"
            f"{cabecalho_str}\n\n"
            f"{nome_grupo}:\n"
            + "\n".join(linhas_formatadas)
        )
        documentos.append(
            Document(
                page_content=conteudo,
                metadata={
                    "source": caminho_arquivo,
                    "sheet": nome_aba,
                    "type": "xlsx",
                    "group": nome_grupo,
                },
            )
        )

    return documentos


def _extrair_texto_url(url: str) -> Document:
    """
    Extrai o conteúdo textual de uma URL removendo agressivamente elementos
    de navegação, rodapés e scripts para reduzir ruído no embedding.
    """
    resposta = requests.get(
        url,
        timeout=30,
        headers={"User-Agent": "Mozilla/5.0 (compatible; assistente-campus/1.0)"},
    )
    resposta.raise_for_status()
    resposta.encoding = resposta.apparent_encoding or resposta.encoding

    sopa = BeautifulSoup(resposta.text, "html.parser")

    # 1. Remover todas as tags de ruído antes de qualquer extração
    for tag in sopa(_TAGS_RUIDO):
        tag.decompose()

    # 2. Tentar extrair apenas o conteúdo principal (article ou main)
    corpo_principal = sopa.find("article") or sopa.find("main") or sopa.find("body") or sopa

    # 3. Extrair texto dos elementos de conteúdo com separador de linha
    blocos = []
    for elemento in corpo_principal.find_all(_TAGS_CONTEUDO):
        texto = elemento.get_text(" ", strip=True)
        # Ignorar blocos muito curtos (menus, rótulos soltos, etc.)
        if len(texto) > 15:
            blocos.append(texto)

    # Deduplicar blocos preservando a ordem (parágrafos repetidos por herança do DOM)
    vistos: set[str] = set()
    blocos_unicos = []
    for bloco in blocos:
        if bloco not in vistos:
            vistos.add(bloco)
            blocos_unicos.append(bloco)

    conteudo = "\n\n".join(blocos_unicos)

    titulo = sopa.title.string.strip() if sopa.title and sopa.title.string else url

    # Prefixar com título e URL para facilitar a citação pelo agente
    conteudo_final = f"Título: {titulo}\nURL: {url}\n\n{conteudo}" if conteudo else f"Título: {titulo}\nURL: {url}"

    print(f"   🌐 URL '{url}' → {len(conteudo_final)} caracteres extraídos")

    return Document(
        page_content=conteudo_final,
        metadata={"source": url, "title": titulo, "type": "url"},
    )


def _carregar_urls_em_arquivo(caminho_arquivo: str) -> list[Document]:
    documentos = []
    with open(caminho_arquivo, "r", encoding="utf-8") as arquivo:
        for linha in arquivo:
            url = linha.strip()
            if not url or url.startswith("#"):
                continue
            try:
                documentos.append(_extrair_texto_url(url))
            except Exception as erro:
                print(f"⚠️ Falha ao carregar URL '{url}': {erro}")
    return documentos


def _carregar_documentos_local(caminho_arquivo: str) -> list[Document]:
    extensao = Path(caminho_arquivo).suffix.lower()

    if extensao == ".pdf":
        return PyPDFLoader(caminho_arquivo).load()
    if extensao == ".txt":
        return [_extrair_texto_txt(caminho_arquivo)]
    if extensao == ".docx":
        return [_extrair_texto_docx(caminho_arquivo)]
    if extensao == ".csv":
        return [_extrair_texto_csv(caminho_arquivo)]
    if extensao == ".xlsx":
        return _extrair_texto_xlsx(caminho_arquivo)
    if extensao == ".url":
        with open(caminho_arquivo, "r", encoding="utf-8") as arquivo:
            url = arquivo.readline().strip()
        if not url:
            return []
        return [_extrair_texto_url(url)]

    return []


def _coletar_documentos() -> list[Document]:
    documentos = []

    if os.path.exists(config.DATA_RAW_DIR):
        for raiz, _, arquivos in os.walk(config.DATA_RAW_DIR):
            for nome_arquivo in arquivos:
                caminho_arquivo = os.path.join(raiz, nome_arquivo)
                extensao = Path(caminho_arquivo).suffix.lower()
                if extensao in SUPPORTED_EXTENSIONS:
                    try:
                        documentos.extend(_carregar_documentos_local(caminho_arquivo))
                    except Exception as erro:
                        print(f"⚠️ Ignorando '{caminho_arquivo}': {erro}")

    if os.path.exists(config.DATA_URLS_FILE):
        try:
            documentos.extend(_carregar_urls_em_arquivo(config.DATA_URLS_FILE))
        except Exception as erro:
            print(f"⚠️ Ignorando URLs em '{config.DATA_URLS_FILE}': {erro}")

    return documentos


def executar_ingestao():
    print("🔄 Iniciando o processo de ingestão de documentos...")

    if not os.path.exists(config.DATA_RAW_DIR) and not os.path.exists(config.DATA_URLS_FILE):
        print(f"❌ Erro: Nenhuma fonte encontrada em '{config.DATA_RAW_DIR}' ou '{config.DATA_URLS_FILE}'.")
        return

    print("📂 Lendo fontes suportadas em data/raw/ e urls.txt...")
    documentos = _coletar_documentos()
    if not documentos:
        print("❌ Erro: Nenhum documento suportado foi encontrado para ingestão.")
        return

    print(f"📄 Total de documentos carregados: {len(documentos)}")

    # Chunks menores para URLs garantem granularidade na busca;
    # overlap maior preserva contexto entre parágrafos fragmentados.
    print("✂️ Dividindo os documentos em blocos de texto (chunks)...")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        add_start_index=True,
        # Separadores em português para quebrar no ponto certo
        separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
    )
    chunks = text_splitter.split_documents(documentos)
    print(f"🧩 Total de chunks gerados: {len(chunks)}")

    # Modelo multilíngue — funciona bem com português
    print("🧠 Gerando embeddings (modelo multilíngue)...")
    embeddings = HuggingFaceEmbeddings(model_name=config.EMBEDDING_MODEL_NAME)

    if os.path.exists(config.CHROMA_DB_DIR):
        try:
            print(f"🧹 Limpando coleção anterior em: {config.CHROMA_DB_DIR}")
            Chroma(
                persist_directory=config.CHROMA_DB_DIR,
                embedding_function=embeddings,
            ).delete_collection()
        except Exception as erro:
            print(f"⚠️ Não foi possível limpar a coleção anterior: {erro}")

    vector_store = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=config.CHROMA_DB_DIR,
    )

    print(f"✅ Ingestão concluída com sucesso! Banco salvo em: {config.CHROMA_DB_DIR}")


if __name__ == "__main__":
    executar_ingestao()